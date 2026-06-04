"""
安评报告智能提取 - Streamlit 入口（含登录与工作台）
运行命令: streamlit run streamlit_app.py
"""
import os
import sys
import tempfile
import json
import io
import uuid
import re
from typing import Optional
from urllib.error import HTTPError, URLError
from urllib.request import Request, urlopen
from pathlib import Path

import streamlit as st

_project_root = Path(__file__).resolve().parent
if str(_project_root) not in sys.path:
    sys.path.insert(0, str(_project_root))

from extract_all_tables import (
    get_all_tables_info,
    filter_tables_for_display,
    extract_all_tables_from_pdf,
)

# ---------------------------------------------------------------------------
# 认证配置
# ---------------------------------------------------------------------------
LOGIN_USER = os.getenv("LOGIN_USERNAME", "admin")
LOGIN_PASS = os.getenv("LOGIN_PASSWORD", "admin123")

st.set_page_config(page_title="安全评估数据平台", page_icon="📊", layout="centered")

# ---------------------------------------------------------------------------
# Session state 初始化
# ---------------------------------------------------------------------------
if "page" not in st.session_state:
    st.session_state.page = "login"
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "username" not in st.session_state:
    st.session_state.username = ""


def do_logout():
    st.session_state.logged_in = False
    st.session_state.username = ""
    st.session_state.page = "login"


def go_workbench():
    st.session_state.page = "workbench"


def go_tool():
    st.session_state.page = "tool"


# ============================ 登录页 ============================
def login_page():
    st.markdown("<br>", unsafe_allow_html=True)
    _, col2, _ = st.columns([1, 2, 1])
    with col2:
        st.markdown("### 安全评估数据平台")
        st.caption("请登录以使用系统功能")
        username = st.text_input("用户名", key="login_username")
        password = st.text_input("密码", type="password", key="login_password")

        def try_login():
            if username == LOGIN_USER and password == LOGIN_PASS:
                st.session_state.logged_in = True
                st.session_state.username = username
                st.session_state.page = "workbench"
            else:
                st.session_state.login_error = True

        if "login_error" in st.session_state and st.session_state.login_error:
            st.error("用户名或密码错误")

        st.button("登 录", key="login_btn", use_container_width=True, type="primary", on_click=try_login)


# ============================ 工作台 ============================
def workbench_page():
    st.markdown("### 工作台")
    st.caption("选择工具开始处理报告数据")
    st.markdown("---")

    col1, col2, col3 = st.columns(3)

    with col1:
        st.markdown(
            """<div style="background:#f0f4ff;border-radius:12px;padding:20px;min-height:220px">
            <h4>安评报告智能提取</h4>
            <p style="color:#666;font-size:14px">上传安全评价报告（PDF/Word），智能识别并提取表格数据，导出为 Word 文档</p>
            </div>""",
            unsafe_allow_html=True,
        )
        st.button("开始使用", key="btn_safety", use_container_width=True, type="primary", on_click=go_tool)

    with col2:
        st.markdown(
            """<div style="background:#f0f4ff;border-radius:12px;padding:20px;min-height:220px;opacity:0.6">
            <h4>环评报告智能提取</h4>
            <p style="color:#666;font-size:14px">上传环境影响评价报告，智能识别并提取关键数据表格</p>
            </div>""",
            unsafe_allow_html=True,
        )
        st.button("开发中", key="btn_env", use_container_width=True, disabled=True)

    with col3:
        st.markdown(
            """<div style="background:#f0f4ff;border-radius:12px;padding:20px;min-height:220px;opacity:0.6">
            <h4>总体规划智能提取</h4>
            <p style="color:#666;font-size:14px">上传总体规划文档，智能识别并提取规划关键信息</p>
            </div>""",
            unsafe_allow_html=True,
        )
        st.button("开发中", key="btn_plan", use_container_width=True, disabled=True)


# ============================ 工具页 ============================
def tool_page():
    st.markdown("### 📄 安评报告智能提取")
    st.caption("支持提取 PDF / Word 中的表格，统一导出为 Word (.docx)")

    # ---- 侧边栏：模式选择 ----
    with st.sidebar:
        st.subheader("运行方式")
        mode = st.radio(
            "选择处理方式",
            options=["本地处理", "连接后端 API"],
            index=0,
        )
        api_base = st.text_input("后端地址（API 模式）", value="http://localhost:5000")
        api_timeout = st.slider("API 超时（秒）", 10, 300, 120, 10)

    uploaded_file = st.file_uploader(
        "上传文件", type=["pdf", "docx"], help="支持 PDF / Word（.docx）"
    )
    if uploaded_file is None:
        return

    filename = (uploaded_file.name or "upload").strip()
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    file_bytes = uploaded_file.getvalue()

    if mode.startswith("本地"):
        _local_process(filename, ext, file_bytes)
    else:
        _api_process(filename, ext, file_bytes, api_base, api_timeout)


# ---------------------------------------------------------------------------
# 本地处理
# ---------------------------------------------------------------------------
def _local_process(filename, ext, file_bytes):
    if ext not in {"pdf", "docx"}:
        st.error("仅支持 PDF 或 Word（.docx）。")
        return

    if ext == "docx":
        with tempfile.TemporaryDirectory() as tmpdir:
            in_path = os.path.join(tmpdir, filename or "upload.docx")
            with open(in_path, "wb") as f:
                f.write(file_bytes)
            try:
                with st.spinner("正在分析 Word 中的表格..."):
                    groups, _ = _get_docx_table_groups(in_path)
            except Exception as e:
                st.error(f"读取 Word 失败：{e}")
                return
            if not groups:
                st.warning("未在该 Word 文档中发现表格。")
                return
            st.success(f"共识别到 **{len(groups)}** 个表格（按表名去重后）。")

            options, option_to_id = [], {}
            for g in groups:
                name = g.get("name") or f"表格"
                cnt = g.get("count") or 1
                label = f"{name}（共 {cnt} 个同名表）" if cnt > 1 else name
                options.append(label)
                option_to_id[label] = g.get("id")

            selected = st.multiselect("选择要保留的表格（不选则保留全部）", options, default=options)
            selected_ids = [option_to_id[o] for o in selected] if selected else None

            if st.button("生成 Word（.docx）", key="gen_docx_word", type="primary"):
                try:
                    with st.spinner("正在生成 Word..."):
                        kept = word_remove_non_table_content(in_path, out_path, selected_ids)
                        if kept <= 0:
                            st.error("文档中无表格或未选择任何表格。")
                            return
                        with open(out_path, "rb") as f:
                            docx_bytes = f.read()
                    st.download_button(
                        "下载结果 Word（.docx）",
                        data=docx_bytes,
                        file_name=f"{Path(filename).stem}_tables_only.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                    st.success(f"已保留 {kept} 个表格。")
                except Exception as e:
                    st.error(f"生成失败：{e}")
    else:
        with tempfile.TemporaryDirectory() as tmpdir:
            pdf_path = os.path.join(tmpdir, filename)
            with open(pdf_path, "wb") as f:
                f.write(file_bytes)
            try:
                with st.spinner("正在识别 PDF 中的表格..."):
                    all_tables = get_all_tables_info(pdf_path)
                    display_tables = filter_tables_for_display(all_tables)
            except Exception as e:
                st.error(f"识别表格失败：{e}")
                return
            if not display_tables:
                st.warning("未在该 PDF 中发现可显示的表格。")
                return
            st.success(f"共识别到 **{len(display_tables)}** 个表格。")

            opts = [f"{t.get('name', t.get('id', ''))}（第{t.get('page', '?')}页）" for t in display_tables]
            opt2id = {opt: t["id"] for t, opt in zip(display_tables, opts)}
            selected = st.multiselect("选择要提取的表格（不选则提取全部）", opts, default=[])
            selected_ids = [opt2id[o] for o in selected] if selected else None

            if st.button("生成 Word（.docx）", key="gen_pdf_word", type="primary"):
                out_dir = os.path.join(tmpdir, "output")
                os.makedirs(out_dir, exist_ok=True)
                try:
                    with st.spinner("正在提取表格并生成 Word..."):
                        result = extract_all_tables_from_pdf(
                            pdf_path, output_dir=out_dir,
                            selected_table_ids=selected_ids, output_format="docx",
                        )
                        tables_data = result.get("tables_data", []) if isinstance(result, dict) else []
                        if not tables_data:
                            st.error("未提取到任何表格数据。")
                            return
                        docx_bytes = _build_docx_from_tables(tables_data)
                    st.download_button(
                        "下载结果 Word（.docx）",
                        data=docx_bytes,
                        file_name=f"{Path(filename).stem}_tables.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                    st.success(f"已提取 {len(tables_data)} 个表格。")
                except Exception as e:
                    st.error(f"提取失败：{e}")


# ---------------------------------------------------------------------------
# 后端 API 处理
# ---------------------------------------------------------------------------
def _api_process(filename, ext, file_bytes, api_base, api_timeout):
    base = api_base.strip()
    if not base:
        st.error("请输入后端地址。")
        return
    try:
        with st.spinner("正在检查后端服务..."):
            health = _http_json("GET", _join_url(base, "/api/health"), timeout=int(api_timeout))
        if not isinstance(health, dict) or health.get("status") != "ok":
            st.warning(f"后端异常：{health}")
    except Exception as e:
        st.error(f"无法连接后端：{e}")
        return
    try:
        with st.spinner("正在上传..."):
            up = _http_upload_file(_join_url(base, "/api/upload"), "file", filename, file_bytes, timeout=int(api_timeout))
        backend_fn = up.get("filename")
        if not backend_fn:
            st.error(f"上传失败：{up}")
            return
    except Exception as e:
        st.error(f"上传失败：{e}")
        return
    try:
        with st.spinner("正在获取表格列表..."):
            resp = _http_json("POST", _join_url(base, "/api/tables"), {"filename": backend_fn}, timeout=int(api_timeout))
        tables = resp.get("tables", []) if isinstance(resp, dict) else []
        if not tables:
            st.warning(f"未获取到表格列表：{resp}")
            return
    except Exception as e:
        st.error(f"获取表格列表失败：{e}")
        return
    st.success(f"共 **{len(tables)}** 个表格。")
    ops = [f"{t.get('name', t.get('id', ''))}（第{t.get('page', '?')}页）" for t in tables]
    o2id = {o: t["id"] for t, o in zip(tables, ops)}
    sel = st.multiselect("选择要提取的表格（不选则提取全部）", ops, default=[])
    selected = [o2id[o] for o in sel] if sel else None
    if st.button("提取并生成 Word", key="api_extract_btn", type="primary"):
        try:
            with st.spinner("后端正在提取..."):
                payload = {"filename": backend_fn}
                if selected is not None:
                    payload["selected_table_ids"] = selected
                result = _http_json("POST", _join_url(base, "/api/extract"), payload=payload, timeout=int(api_timeout))
            dl = result.get("download_url")
            out_name = result.get("output_filename") or "result.docx"
            if not dl:
                st.error(f"提取失败：{result}")
                return
            with st.spinner("正在下载..."):
                out_bytes = _http_get_bytes(_join_url(base, dl), timeout=int(api_timeout))
            st.download_button(
                "下载结果文件", data=out_bytes, file_name=out_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            st.success(f"已生成：{out_name}（表格数量：{result.get('total_tables', '未知')}）")
        except Exception as e:
            st.error(f"提取失败：{e}")


# ---------------------------------------------------------------------------
# 共享工具函数（原版保留）
# ---------------------------------------------------------------------------
def _join_url(base: str, path: str) -> str:
    base = (base or "").rstrip("/")
    if not base:
        return path
    if not path.startswith("/"):
        path = "/" + path
    return base + path


def _http_json(method: str, url: str, payload: dict | None = None, timeout: int = 60) -> dict:
    data = None
    headers = {"Accept": "application/json"}
    if payload is not None:
        data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        headers["Content-Type"] = "application/json; charset=utf-8"
    req = Request(url, data=data, headers=headers, method=method.upper())
    with urlopen(req, timeout=timeout) as resp:
        raw = resp.read()
        return json.loads(raw.decode("utf-8")) if raw else {}


def _encode_multipart(field: str, fname: str, content: bytes, ct: str = "application/octet-stream") -> tuple:
    boundary = uuid.uuid4().hex
    body = io.BytesIO()
    body.write(f"--{boundary}\r\n".encode())
    body.write(f'Content-Disposition: form-data; name="{field}"; filename="{fname}"\r\nContent-Type: {ct}\r\n\r\n'.encode())
    body.write(content)
    body.write(b"\r\n")
    body.write(f"--{boundary}--\r\n".encode())
    return body.getvalue(), f"multipart/form-data; boundary={boundary}"


def _http_upload_file(url: str, field: str, fname: str, content: bytes, timeout: int = 120) -> dict:
    body, ct = _encode_multipart(field, fname, content)
    req = Request(url, data=body, headers={"Content-Type": ct, "Accept": "application/json"}, method="POST")
    with urlopen(req, timeout=timeout) as resp:
        raw = resp.read()
        return json.loads(raw.decode("utf-8")) if raw else {}


def _http_get_bytes(url: str, timeout: int = 120) -> bytes:
    with urlopen(Request(url, method="GET"), timeout=timeout) as resp:
        return resp.read()


def _clean_cell_text(text: str) -> str:
    if not text or not isinstance(text, str):
        return ""
    s = text.replace("\r\n", "\n").replace("\r", "\n").replace("\a", "\n")
    s = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", s)
    return s.strip()


def _normalize_docx_table_title(pt: Optional[str]) -> Optional[str]:
    if not pt or not isinstance(pt, str):
        return None
    s = pt.strip()
    if not s:
        return None
    if "评价报告摘要" in s:
        return "评价报告摘要"
    m = re.match(r"^表\s*\d+\s*[-－–]\s*\d+\s+.+", s)
    if m:
        return s
    m = re.search(r"表\s*\d+\s*[-－–]\s*\d+\s+.+", s)
    return m.group(0).strip() if m else None


def _clean_docx_title_line(pt: Optional[str]) -> Optional[str]:
    if not pt or not isinstance(pt, str):
        return None
    s = _clean_cell_text(pt)
    if not s:
        return None
    for line in s.split("\n"):
        t = re.sub(r"\s+", " ", line.strip())
        if t:
            return t
    return None


def _dedupe_title_text(title: str) -> str:
    if not title or not isinstance(title, str):
        return ""
    s = re.sub(r"\s+", " ", title).strip()
    m = re.match(r"^(.+?)(?:\s+\1)+$", s)
    if m:
        s = m.group(1).strip()
    tokens = s.split(" ")
    out, prev = [], None
    for tok in tokens:
        if tok and tok != prev:
            out.append(tok)
        prev = tok
    if len(out) % 2 == 0 and out[: len(out) // 2] == out[len(out) // 2 :]:
        out = out[: len(out) // 2]
    return " ".join(out).strip()


def _get_docx_table_groups(docx_path: str):
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(docx_path)
    body = doc.element.body
    p_tag = qn("w:p")
    tbl_tag = qn("w:tbl")

    last_text = None
    groups_by_name = {}
    groups_in_order = []
    ti = 0

    for child in body:
        if child.tag == p_tag:
            try:
                t = "".join((n.text or "") for n in child.iter() if hasattr(n, "text"))
                if t and t.strip():
                    last_text = t.strip()
            except Exception:
                pass
            continue
        if child.tag != tbl_tag:
            continue
        tid = f"table_{ti}"
        raw = _clean_docx_title_line(last_text)
        norm = _normalize_docx_table_title(raw) if raw else None
        name = _dedupe_title_text(norm or raw or f"表格{ti + 1}")
        if name not in groups_by_name:
            grp = {
                "id": tid, "name": name, "page": ti + 1,
                "table_num": len(groups_in_order) + 1,
                "table_ids": [tid], "count": 1,
            }
            groups_by_name[name] = grp
            groups_in_order.append(grp)
        else:
            grp = groups_by_name[name]
            grp["table_ids"].append(tid)
            grp["count"] += 1
        ti += 1

    id_map = {}
    for g in groups_in_order:
        for tid in g["table_ids"]:
            id_map[tid] = g["table_ids"]
        id_map[g["id"]] = g["table_ids"]
    return groups_in_order, id_map


def word_remove_non_table_content(docx_path, out_path, selected_ids=None):
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(docx_path)
    body = doc.element.body
    tbl_tag = qn("w:tbl")
    p_tag = qn("w:p")
    children = list(body)
    keep = set()
    ti = 0

    for i, child in enumerate(children):
        if child.tag == tbl_tag:
            if selected_ids is None or f"table_{ti}" in selected_ids:
                keep.add(i)
                if i > 0 and children[i - 1].tag == p_tag:
                    keep.add(i - 1)
            ti += 1

    for i in range(len(children) - 1, -1, -1):
        if i not in keep:
            body.remove(children[i])
    doc.save(out_path)
    return sum(1 for i in keep if children[i].tag == tbl_tag)


def _build_docx_from_tables(tables_data):
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "等线"
    style.font.size = Pt(9)
    try:
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    except Exception:
        pass

    for t in tables_data:
        title = (t.get("title") or t.get("name") or "").strip()
        if title:
            p = doc.add_paragraph()
            r = p.add_run(title)
            r.bold = True
            r.font.size = Pt(10.5)
        data = t.get("data") or []
        if not data:
            continue
        ncols = max((len(row) for row in data), default=0)
        if ncols <= 0:
            continue
        norm = []
        for row in data:
            row = list(row or [])
            while len(row) < ncols:
                row.append("")
            norm.append(row[:ncols])
        tbl = doc.add_table(rows=len(norm), cols=ncols, style="Table Grid")
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ri, row in enumerate(norm):
            for ci, val in enumerate(row):
                tbl.cell(ri, ci).text = "" if val is None else str(val)
        doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ============================ 主流程 ============================
if st.session_state.page == "login" or not st.session_state.logged_in:
    login_page()
elif st.session_state.page == "workbench":
    # 顶部栏
    col_l, col_r = st.columns([6, 1])
    with col_l:
        st.caption("安全评估数据平台")
    with col_r:
        st.button("退出", key="logout_btn", use_container_width=True, on_click=do_logout)
    workbench_page()
else:
    # tool 页
    col_l, col_r = st.columns([6, 1])
    with col_l:
        st.caption("安评报告智能提取")
    with col_r:
        st.button("退出", key="logout_btn2", use_container_width=True, on_click=do_logout)
    col_back, _ = st.columns([1, 6])
    with col_back:
        st.button("← 返回工作台", key="back_btn", on_click=go_workbench)
    tool_page()
